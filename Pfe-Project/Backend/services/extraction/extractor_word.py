import docx
import os
import zipfile
import easyocr

print("Loading EasyOCR Model for Word documents...")
reader = easyocr.Reader(['en', 'fr'], gpu=False) 

def extract_all_from_docx(file_path: str) -> str:
    """
    Extracts standard text, table text, AND runs OCR on any embedded images 
    inside a Microsoft Word (.docx) file.
    """
    if not os.path.exists(file_path):
        return "Error: File does not exist."

    full_text = []

    try:
        # PART 1: EXTRACT STANDARD TEXT & TABLES
        doc = docx.Document(file_path)
        
        for para in doc.paragraphs:
            clean_text = para.text.strip()
            if clean_text:
                full_text.append(clean_text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    clean_cell_text = cell.text.strip()
                    if clean_cell_text and clean_cell_text not in full_text:
                        full_text.append(clean_cell_text)

        # PART 2: EXTRACT IMAGES & RUN OCR
        with zipfile.ZipFile(file_path, "r") as docx_zip:
            for item in docx_zip.namelist():
                if item.startswith("word/media/") and item.lower().endswith(('.png', '.jpg', '.jpeg')):
                    print(f"Found embedded image: {item}. Running EasyOCR...")
                    
                    image_bytes = docx_zip.read(item)
                    ocr_results = reader.readtext(image_bytes, detail=0)
                    
                    if ocr_results:
                        full_text.append("\n--- [Text Extracted from Image] ---")
                        full_text.extend(ocr_results) 
                        full_text.append("-----------------------------------\n")

        return "\n".join(full_text)

    except Exception as e:
        return f"An error occurred while reading the DOCX file: {e}"